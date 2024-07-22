import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from pathlib import Path

# Load the Excel file
excel_file = Path(r'C:\Users\sau23\Downloads\modified_Applicant_Observation_final_V.1.0.docx')
# df = pd.read_excel(excel_file)
df = oc = Document(excel_file)

# Print column names to troubleshoot
print("Column names:", df.columns.tolist())

# Strip whitespace from column names
df.columns = df.columns.str.strip()

# Create a new Document
doc = Document()

# Define a function to add formatted text
def add_formatted_paragraph(doc, title, text, title_format, text_format):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.name = title_format['font_name']
    run.font.size = Pt(title_format['font_size'])
    run.font.bold = title_format['bold']
    run.font.color.rgb = RGBColor(*title_format['color'])
    
    run = p.add_run(text)
    run.font.name = text_format['font_name']
    run.font.size = Pt(text_format['font_size'])
    run.font.bold = text_format['bold']
    run.font.color.rgb = RGBColor(*text_format['color'])

# Define text formats
formats = {
    'title': {'font_name': 'Calibri', 'font_size': 14, 'bold': True, 'color': (15, 71, 97)},
    'text': {'font_name': 'Calibri', 'font_size': 14, 'bold': True, 'color': (0, 112, 192)},
    'section_title': {'font_name': 'Calibri', 'font_size': 12, 'bold': False, 'color': (0, 0, 0)},
    'section_text': {'font_name': 'Calibri', 'font_size': 12, 'bold': True, 'color': (0, 0, 0)},
    'cve_text': {'font_name': 'Calibri', 'font_size': 12, 'bold': True, 'color': (0, 112, 192)},
    'severity_high': {'font_name': 'Calibri', 'font_size': 12, 'bold': True, 'color': (255, 0, 0)},
    'severity_critical': {'font_name': 'Calibri', 'font_size': 12, 'bold': True, 'color': (192, 0, 0)},
    'severity_medium': {'font_name': 'Calibri', 'font_size': 12, 'bold': True, 'color': (255, 192, 0)},
    'severity_low': {'font_name': 'Calibri', 'font_size': 12, 'bold': True, 'color': (0, 176, 80)}
}

# Process each row in the dataframe
for index, row in df.iterrows():
    # add_formatted_paragraph(doc, 'Vulnerability Reference No: ', str(row['Vulnerability Unique ID']), formats['title'], formats['text'])
    add_formatted_paragraph(doc, 'Vulnerability Reference No: ', str(row['Sl.No.']), formats['title'], formats['text'])
    
    # # add_formatted_paragraph(doc, 'I. Affected Asset i.e. IP/URL/Application etc.: ', str(row['Affected Asset i.e., IP/URL/Application etc']), formats['section_title'], formats['section_text'])
    # add_formatted_paragraph(doc, 'I. Affected Asset i.e. IP/URL/Application etc.: ', str(row['Affected Assets i.e.,IP/URL/Application etc.']), formats['section_title'], formats['section_text'])
    # add_formatted_paragraph(doc, 'II. Observation/ Vulnerability title: ', str(row['Observation\Vulnerability Title']), formats['section_title'], formats['section_text'])
    # add_formatted_paragraph(doc, 'III. Detailed observation / Vulnerable point: ', str(row['Detailed observation / Vulnerable point']), formats['section_title'], formats['section_text'])
    # add_formatted_paragraph(doc, 'IV. CVE/CWE: ', str(row['CVE/CWE']), formats['section_title'], formats['cve_text'])
    
    # # Setting Control Objective, Control Name, and Audit Requirement to NA
    # add_formatted_paragraph(doc, 'V. Control Objective: ', 'NA', formats['section_title'], formats['section_text'])
    # add_formatted_paragraph(doc, 'VI. Control Name: ', 'NA', formats['section_title'], formats['section_text'])
    # add_formatted_paragraph(doc, 'VII. Audit Requirement: ', 'NA', formats['section_title'], formats['section_text'])
    
    # severity = str(row['Severity']).lower()
    # if severity == 'high':
    #     severity_format = formats['severity_high']
    # elif severity == 'critical':
    #     severity_format = formats['severity_critical']
    # elif severity == 'medium':
    #     severity_format = formats['severity_medium']
    # elif severity == 'low':
    #     severity_format = formats['severity_low']
    # else:
    #     severity_format = formats['section_text']

    # add_formatted_paragraph(doc, 'VIII. Severity: ', str(row['Severity']), formats['section_title'], severity_format)
    
    # add_formatted_paragraph(doc, 'IX. Recommendation: ', str(row['Recommendation']), formats['section_title'], formats['section_text'])
    # add_formatted_paragraph(doc, 'X. Reference: ', str(row['Reference']), formats['section_title'], formats['cve_text'])
    # add_formatted_paragraph(doc, 'XI. New or Repeat observation: ', str(row['New or Repeat Observation']), formats['section_title'], formats['section_text'])
    # add_formatted_paragraph(doc, 'XII. References to evidences / Proof of Concept: ', str(row['POC']), formats['section_title'], formats['section_text'])
    
    # doc.add_page_break()
    serial_num = 1
    for para in doc.paragraphs:
        # Check if the paragraph contains the placeholder text to be replaced
        prefix = "Vulnerability Reference No: VULN-Applicant-"
        
        if prefix in para.text:
            # Generate the new serial number with leading zeros
            serial_number = f"{prefix}{str(serial_num).zfill(3)}"
            # Replace the placeholder text with the generated serial number
            para.text = para.text.replace(para.text, serial_number)
            # Increment the serial number for the next replacement
            serial_num += 1

# Save the document
doc.save('report2.docx')
