from docx import Document
import pandas as pd

def read_docx(file_path):
    doc = Document(file_path)
    full_text = []
    vuln_text = []
    recommendation = []
    severity = []
    reference = []
    # title = "II. Observation/ Vulnerability Title: "
    title = "II. Observation/ Vulnerability title:"
    # vuln_name = "III. Detailed observation / Vulnerable point: "
    vuln_name = "III. Detailed observation / Vulnerable point:"
    end_keyword = "IV. CVE/CWE:"
    sever = "VIII. Severity:"
    recommend = "IX. Recommendation:"
    ref = "X. Reference:"
    cv = "IV. CVE/CWE:"
    obj = "V. Control Objective:"
    cve_text = []
    is_between_keywords = False
    is_between_keywords1 = False
    is_between_keywords2 = False
    test = []
    test1 =[]
    test2 = []
    serial = "Vulnerability Reference No:"
    serial_num = []
    for para in doc.paragraphs:
        
        if serial in para.text:
            serial_num.append(para.text[len(serial):].strip())
            
        if title in para.text:
            full_text.append(para.text[len(title):].strip())
            # print(para.text[len(title):])
        # for para in doc.paragraphs:
        # test = []
        if vuln_name in para.text:
            is_between_keywords = True


        if is_between_keywords:
            if end_keyword in para.text:
                    is_between_keywords = False
                    # vuln_text = ["".join(vuln_text)]
                    # break
                    
            else:
                if vuln_name in para.text:
                    test.append(para.text[len(vuln_name):])
                else:
                    test.append(para.text)
                    
        if not is_between_keywords and len(test) !=0:
            vuln_text.append("".join(test))
            test = []
        
        
        ### CVE pull
        if cv in para.text:
            is_between_keywords2 = True

        if is_between_keywords2:
            if obj in para.text:
                    is_between_keywords2 = False
                    # vuln_text = ["".join(vuln_text)]
                    # break
                    
            else:
                if cv in para.text:
                    test2.append(para.text[len(cv):])
                else:
                    test2.append(para.text)
                    
        if not is_between_keywords2 and len(test2) !=0:
            cve_text.append("".join(test2))
            test2 = []

        if sever in para.text:
            severity.append(para.text[len(sever):])
         
        if recommend in para.text:
            is_between_keywords1 = True
               
        if is_between_keywords1:
            if ref in para.text:
                    is_between_keywords1 = False
            else:
                if recommend in para.text:
                    test1.append(para.text[len(recommend):])
                else:
                    test1.append(para.text)
                    
        if not is_between_keywords1 and len(test1) !=0:
            recommendation.append("".join(test1))
            test1 = []
            
            
        if ref in para.text:
            reference.append(para.text[len(ref):])
        
    
    # print(len(vuln_text), len(full_text), len(recommendation), len(reference), len(severity))
    # serial_numbers = [f"VULN-Applicant-{str(i).zfill(3)}" for i in range(1, len(full_text) + 1)]
    data = {
        "Sl.No.": serial_num,
        "Affected Assets i.e.,IP/URL/Application etc.": "192.168.125.18",
    "Observation/ Vulnerability Title": full_text,
    "CVE/CWE": cve_text,
    "Severity": severity,
    "Control Objective#": "N/A",
    "Control Name#": "N/A",
    "Audit Requirement": "N/A",
    "Recommendation": recommendation,
    "Reference": reference,
    "New or Repeat Observation": "New Observation",
    "Detailed observation / Vulnerable point": vuln_text,
    "POC": ""
}
    print(len(full_text), len(vuln_text), len(recommendation), len(reference), len(severity))
    # print(full_text, vuln_text, recommendation, reference, severity)
    final_data = pd.DataFrame(data)
    
    final_data.to_excel("C:\\Users\\sau23\\Desktop\\modified_Applicant_Observation_final_V.3.0.xlsx",index=False)
    return 1 #'\n'.join(full_text)


# Example usage:
if __name__ == "__main__":
    docx_file = 'C:\\Users\\sau23\\Downloads\\modified_Applicant_Observation_final_V.3.0.docx'
    extracted_text = read_docx(docx_file)
    print(extracted_text)
